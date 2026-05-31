import os
import io
import re
import json
import base64
import traceback
import telebot
from groq import Groq
from flask import Flask, request
from duckduckgo_search import DDGS
from bs4 import BeautifulSoup

BOT_TOKEN = os.environ.get('BOT_TOKEN')
GROQ_API_KEY = os.environ.get('GROQ_API_KEY')
NVIDIA_API_KEY = os.environ.get('NVIDIA_API_KEY')
OPENROUTER_API_KEY = os.environ.get('OPENROUTER_API_KEY')

app = Flask(__name__)
bot = telebot.TeleBot(BOT_TOKEN, threaded=False)
groq_client = Groq(api_key=GROQ_API_KEY)

current_model = "groq"
NVIDIA_BASE = "https://integrate.api.nvidia.com/v1"
OPENROUTER_BASE = "https://openrouter.ai/api/v1"

SEARCH_TRIGGERS = ['cari', 'search', 'google', 'latest', 'terbaru', 'berita', 'news', 'update', 'terkini', 'hari ini']

def nvidia_headers():
    return {"Authorization": f"Bearer {NVIDIA_API_KEY}", "Content-Type": "application/json"}

def extract_url(text):
    urls = re.findall(r'https?://[^\s<>"]+|www\.[^\s<>"]+', text)
    return urls[0] if urls else None

def read_webpage(url):
    import requests
    try:
        resp = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=30)
        if resp.status_code != 200:
            return None, f"HTTP {resp.status_code}"
        soup = BeautifulSoup(resp.text, 'lxml')
        for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
            tag.decompose()
        title = soup.title.string.strip() if soup.title and soup.title.string else ""
        text = soup.get_text(separator='\n', strip=True)
        lines = [l for l in text.split('\n') if len(l) > 20]
        content = "\n".join(lines[:500])
        if len(content) > 15000:
            content = content[:15000] + "..."
        return content, title
    except Exception as e:
        return None, str(e)

def web_search(query, max_results=5):
    try:
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=max_results))
            return results
    except Exception as e:
        print(f"Search error: {e}", flush=True)
        return []

def should_search(text):
    t = text.lower().strip()
    if t.startswith('/search') or t.startswith('!cari') or t.startswith('cari '):
        return True
    return any(trigger in t for trigger in SEARCH_TRIGGERS)

def should_generate(text):
    if text.startswith('/generate'):
        return True
    return any(text.lower().startswith(g) for g in ['gambar ', 'buat ', 'buatkan ', 'gambar:', 'buat:'])

def format_results(results):
    if not results:
        return ""
    out = "\n\n--- HASIL PENCARIAN WEB ---\n"
    for i, r in enumerate(results[:5], 1):
        out += f"{i}. {r.get('title', '')}\n{r.get('body', '')}\n{r.get('href', '')}\n\n"
    return out

def ask_groq(prompt, max_tokens=1500):
    resp = groq_client.chat.completions.create(
        model="llama-3.3-70b-versatile", messages=[{"role": "user", "content": prompt}], max_tokens=max_tokens)
    return resp.choices[0].message.content

def ask_nvidia(prompt, max_tokens=1500):
    import requests
    resp = requests.post(f"{NVIDIA_BASE}/chat/completions", headers=nvidia_headers(), json={
        "model": "meta/llama-3.3-70b-instruct",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": max_tokens, "temperature": 0.7}, timeout=30)
    if resp.status_code != 200:
        raise Exception(f"NVIDIA {resp.status_code}: {resp.text[:200]}")
    return resp.json()['choices'][0]['message']['content']

def ask_nvidia_vision(prompt, data_url):
    import requests
    resp = requests.post(f"{NVIDIA_BASE}/chat/completions", headers=nvidia_headers(), json={
        "model": "nvidia/llama-3.2-90b-vision-nim",
        "messages": [{"role": "user", "content": [
            {"type": "text", "text": prompt},
            {"type": "image_url", "image_url": {"url": data_url}}]}],
        "max_tokens": 2000}, timeout=30)
    if resp.status_code != 200:
        raise Exception(f"NVIDIA Vision {resp.status_code}: {resp.text[:200]}")
    return resp.json()['choices'][0]['message']['content']

def ask_openrouter(prompt, max_tokens=1500):
    import requests
    resp = requests.post(f"{OPENROUTER_BASE}/chat/completions", headers={
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://telegram-ai-bot-tau.vercel.app",
        "X-Title": "Fredi Telegram AI Bot"
    }, json={
        "model": "openrouter/free",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": max_tokens}, timeout=30)
    if resp.status_code != 200:
        raise Exception(f"OpenRouter {resp.status_code}: {resp.text[:200]}")
    return resp.json()['choices'][0]['message']['content']

def ask_openrouter_vision(prompt, data_url):
    import requests
    for model in ["nvidia/nemotron-nano-12b-2-vl:free", "openrouter/free"]:
        try:
            resp = requests.post(f"{OPENROUTER_BASE}/chat/completions", headers={
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type": "application/json",
                "HTTP-Referer": "https://telegram-ai-bot-tau.vercel.app",
                "X-Title": "Fredi Telegram AI Bot"
            }, json={
                "model": model,
                "messages": [{"role": "user", "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": data_url}}]}],
                "max_tokens": 2000}, timeout=30)
            if resp.status_code == 200:
                return resp.json()['choices'][0]['message']['content']
            err = resp.text.lower()
            if "does not support image" in err or "vision" in err:
                continue
            raise Exception(f"OpenRouter Vision {resp.status_code}: {resp.text[:200]}")
        except Exception as e:
            if "does not support image" in str(e).lower() or "vision" in str(e).lower():
                continue
            raise
    raise Exception("Semua model OpenRouter gagal untuk vision")

def generate_image(prompt):
    import requests
    resp = requests.post(f"{NVIDIA_BASE}/images/generations", headers=nvidia_headers(), json={
        "model": "stabilityai/stable-diffusion-3.5-large",
        "prompt": prompt, "width": 1024, "height": 1024}, timeout=60)
    if resp.status_code != 200:
        raise Exception(f"NVIDIA Image {resp.status_code}: {resp.text[:200]}")
    b64 = resp.json()['data'][0]['b64_json']
    buf = io.BytesIO(base64.b64decode(b64))
    buf.seek(0)
    return buf

def generate_pdf(text):
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=9)
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            pdf.ln(3)
        else:
            safe = line.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 5, safe)
    buf = io.BytesIO(pdf.output())
    buf.seek(0)
    return buf

def generate_docx(text):
    from docx import Document
    doc = Document()
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            pass
        elif line.startswith('#') or line.startswith('**'):
            doc.add_heading(line.replace('#', '').replace('*', '').strip(), level=2)
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

MAX_MSG = 4000

def send_long(chat_id, text, reply_to=None):
    if len(text) <= MAX_MSG:
        return bot.send_message(chat_id, text, reply_to_message_id=reply_to)
    chunks = []
    while text:
        if len(text) <= MAX_MSG:
            chunks.append(text)
            break
        split = text.rfind('\n\n', 0, MAX_MSG)
        if split == -1:
            split = text.rfind('\n', 0, MAX_MSG)
        if split == -1:
            split = text.rfind('. ', 0, MAX_MSG)
        if split == -1 or split < MAX_MSG // 2:
            split = MAX_MSG
        else:
            split += 1
        chunks.append(text[:split].strip())
        text = text[split:].strip()
    first = True
    for chunk in chunks:
        if first:
            bot.send_message(chat_id, chunk, reply_to_message_id=reply_to)
            first = False
        else:
            bot.send_message(chat_id, f"⏩ Lanjutan:\n{chunk}")
    return chunks[0] if chunks else None

def send_menu(chat_id, text):
    from telebot.types import ReplyKeyboardMarkup, KeyboardButton
    markup = ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    markup.add(
        KeyboardButton("/generate"), KeyboardButton("/search"),
        KeyboardButton("/model"), KeyboardButton("/help"),
        KeyboardButton("/setmodel groq"), KeyboardButton("/setmodel nvidia"),
        KeyboardButton("/setmodel openrouter"),
    )
    bot.send_message(chat_id, text, reply_markup=markup)

MENU_TEXT = (
    "📋 MENU BOT\n\n"
    "💬 Tanya apa saja — AI menjawab\n"
    "🎨 /generate <deskripsi> — buat gambar\n"
    "🌐 /search <query> atau 'cari ...' — info terbaru\n"
    "🔗 Tempel link web — AI baca soal/artikel\n"
    "🖼️ Kirim gambar — AI baca & analisis teks\n"
    "⚙️ /setmodel groq — Groq\n"
    "⚙️ /setmodel nvidia — NVIDIA NIM\n"
    "⚙️ /setmodel openrouter — OpenRouter (free models)\n"
    "📌 Contoh: tempel https://...soal-matematika"
)

@bot.message_handler(commands=['start', 'menu', 'help'])
def start(message):
    names = {"groq": "Groq", "nvidia": "NVIDIA NIM", "openrouter": "OpenRouter (Free)"}
    provider = names.get(current_model, current_model)
    send_menu(message.chat.id, f"🤖 Asisten AI Fredi — AI: {provider}\n\n{MENU_TEXT}")

@bot.message_handler(commands=['model'])
def model_info(message):
    info = {
        "groq": ("Groq (Llama 3.3 70B)", "Groq Vision"),
        "nvidia": ("NVIDIA NIM (Llama 3.3 70B)", "NVIDIA Vision"),
        "openrouter": ("OpenRouter (Free Models)", "OpenRouter Vision"),
    }
    p, img = info.get(current_model, ("Unknown", "Unknown"))
    send_menu(message.chat.id, f"🤖 AI: {p}\n🖼️ Gambar: {img}\n🎨 Generate: Stable Diffusion 3.5\n\nGunakan /help")

@bot.message_handler(commands=['setmodel'])
def set_model(message):
    global current_model
    c = message.text.replace('/setmodel', '', 1).strip().lower()
    if c == 'groq':
        current_model = 'groq'
        send_menu(message.chat.id, "✅ Beralih ke Groq")
    elif c in ('nvidia', 'nim'):
        if not NVIDIA_API_KEY:
            send_menu(message.chat.id, "❌ NVIDIA_API_KEY belum diset di Vercel")
            return
        current_model = 'nvidia'
        send_menu(message.chat.id, "✅ Beralih ke NVIDIA NIM (Llama 3.3 + SD3.5)")
    elif c in ('openrouter', 'or'):
        if not OPENROUTER_API_KEY:
            send_menu(message.chat.id, "❌ OPENROUTER_API_KEY belum diset di Vercel")
            return
        current_model = 'openrouter'
        send_menu(message.chat.id, "✅ Beralih ke OpenRouter (Free Models: DeepSeek V4, Gemma, dll)")
    else:
        send_menu(message.chat.id, "Gunakan: /setmodel groq / /setmodel nvidia / /setmodel openrouter")

@bot.message_handler(commands=['generate'])
def generate_command(message):
    prompt = message.text.replace('/generate', '', 1).strip()
    if not prompt:
        send_menu(message.chat.id, "Gunakan: /generate <deskripsi>")
        return
    bot.send_chat_action(message.chat.id, 'upload_photo')
    msg = bot.reply_to(message, "🎨 Membuat gambar dengan Stable Diffusion 3.5...")
    try:
        buf = generate_image(prompt)
        bot.delete_message(message.chat.id, msg.message_id)
        bot.send_photo(message.chat.id, buf, caption=f"🎨 {prompt}")
    except Exception as e:
        bot.edit_message_text(f"Gagal: {str(e)}", message.chat.id, msg.message_id)

@bot.message_handler(commands=['search'])
def search_command(message):
    q = message.text.replace('/search', '', 1).strip()
    if not q:
        send_menu(message.chat.id, "Gunakan: /search <query>")
        return
    bot.send_chat_action(message.chat.id, 'typing')
    results = web_search(q)
    ctx = format_results(results)
    prompt = f"Pertanyaan: {q}\n{ctx}\nJawab berdasarkan hasil pencarian di atas."
    p = current_model
    try:
        if p == 'groq':
            reply = ask_groq(prompt)
        elif p == 'nvidia':
            reply = ask_nvidia(prompt)
        else:
            reply = ask_openrouter(prompt)
        send_long(message.chat.id, reply, reply_to=message.message_id)
    except Exception as e:
        try:
            if p == 'groq':
                reply = ask_nvidia(prompt)
            elif p == 'nvidia':
                reply = ask_groq(prompt)
            else:
                reply = ask_groq(prompt)
            send_long(message.chat.id, f"[Fallback]\n\n{reply}", reply_to=message.message_id)
        except:
            bot.reply_to(message, f"Error: {str(e)}")

@bot.message_handler(commands=['savepdf'])
def save_pdf(message):
    if not message.reply_to_message or not message.reply_to_message.text:
        send_menu(message.chat.id, "Balas pesan dengan /savepdf")
        return
    try:
        buf = generate_pdf(message.reply_to_message.text)
        bot.send_document(message.chat.id, buf, visible_file_name="output.pdf", caption="📄 PDF siap!")
    except Exception as e:
        bot.reply_to(message, f"Gagal: {str(e)}")

@bot.message_handler(commands=['savedocx'])
def save_docx(message):
    if not message.reply_to_message or not message.reply_to_message.text:
        send_menu(message.chat.id, "Balas pesan dengan /savedocx")
        return
    try:
        buf = generate_docx(message.reply_to_message.text)
        bot.send_document(message.chat.id, buf, visible_file_name="output.docx", caption="📝 Word siap!")
    except Exception as e:
        bot.reply_to(message, f"Gagal: {str(e)}")

@bot.message_handler(content_types=['photo'])
def handle_photo(message):
    try:
        bot.send_chat_action(message.chat.id, 'typing')
        file_id = message.photo[-1].file_id; file_info = bot.get_file(file_id)
        downloaded = bot.download_file(file_info.file_path)
        b64 = base64.b64encode(downloaded).decode('utf-8')
        data_url = f"data:image/jpeg;base64,{b64}"
        prompt = (message.caption or "").strip() or "Baca dan analisis semua teks yang ada di gambar ini."
        p = current_model
        if p == 'groq':
            reply = groq_client.chat.completions.create(
                model="llama-3.2-90b-vision-preview",
                messages=[{"role": "user", "content": [{"type": "text", "text": prompt}, {"type": "image_url", "image_url": {"url": data_url}}]}],
                max_tokens=2000).choices[0].message.content
        elif p == 'nvidia':
            reply = ask_nvidia_vision(prompt, data_url)
        else:
            reply = ask_openrouter_vision(prompt, data_url)
        bot.reply_to(message, reply)
    except Exception as e:
        try:
            if current_model == 'nvidia':
                reply = groq_client.chat.completions.create(
                    model="llama-3.2-90b-vision-preview",
                    messages=[{"role": "user", "content": [{"type": "text", "text": prompt}, {"type": "image_url", "image_url": {"url": data_url}}]}],
                    max_tokens=2000).choices[0].message.content
            else:
                reply = ask_nvidia_vision(prompt, data_url)
            bot.reply_to(message, f"[Fallback] {reply}")
        except Exception as e2:
            bot.reply_to(message, f"Maaf, gagal: {str(e2)}")

@bot.message_handler(func=lambda m: True)
def handle(message):
    text = message.text
    bot.send_chat_action(message.chat.id, 'typing')
    if should_generate(text):
        prompt = text
        for p in ['/generate', 'gambar ', 'buat ', 'buatkan ', 'gambar:', 'buat:']:
            if text.lower().startswith(p):
                prompt = text[len(p):].strip()
                break
        if prompt:
            msg = bot.reply_to(message, "🎨 Membuat gambar dengan SD3.5...")
            try:
                buf = generate_image(prompt); bot.delete_message(message.chat.id, msg.message_id)
                bot.send_photo(message.chat.id, buf, caption=f"🎨 {prompt}")
            except Exception as e:
                bot.edit_message_text(f"Gagal: {str(e)}", message.chat.id, msg.message_id)
            return
    url = extract_url(text)
    if url:
        msg = bot.reply_to(message, "🔗 Membaca halaman web...")
        content, title = read_webpage(url)
        if content is None:
            bot.edit_message_text(f"Gagal: {title}", message.chat.id, msg.message_id); return
        question = text.replace(url, '').strip()
        prompt = f"Konten: {title}\n\n{content}\n\n"
        prompt += f"Pertanyaan: {question}\n\nJawab berdasarkan konten di atas secara lengkap." if question else "Baca SELURUH konten di atas dengan teliti. Jika ada soal-soal, tulis dan jawab SEMUA soal satu per satu tanpa ada yang terlewat."
        bot.edit_message_text("🔗 Menganalisis...", message.chat.id, msg.message_id)
        try:
            p = current_model
            if p == 'groq':
                reply = ask_groq(prompt, 4096)
            elif p == 'nvidia':
                reply = ask_nvidia(prompt, 4096)
            else:
                reply = ask_openrouter(prompt, 4096)
            save_note = "\n\n—\n💾 Balas pesan ini dengan /savepdf atau /savedocx untuk simpan sebagai file."
            full = reply + save_note
            if len(full) <= MAX_MSG:
                bot.edit_message_text(full, message.chat.id, msg.message_id)
            else:
                bot.edit_message_text(reply[:MAX_MSG] + "\n\n✂️ Bersambung...", message.chat.id, msg.message_id)
                send_long(message.chat.id, "⏩ Lanjutan:\n" + reply[MAX_MSG:] + save_note)
        except Exception as e:
            bot.edit_message_text(f"Error: {str(e)}", message.chat.id, msg.message_id)
        return
    if should_search(text):
        q = text
        for pfx in ['/search', '!cari', 'cari ']:
            if text.lower().startswith(pfx):
                q = text[len(pfx):].strip(); break
        results = web_search(q); ctx = format_results(results)
        prompt = f"Pertanyaan: {q}\n{ctx}\nJawab berdasarkan hasil pencarian di atas."
    else:
        prompt = text
    p = current_model
    try:
        if p == 'groq':
            reply = ask_groq(prompt)
        elif p == 'nvidia':
            reply = ask_nvidia(prompt)
        else:
            reply = ask_openrouter(prompt)
        send_long(message.chat.id, reply, reply_to=message.message_id)
    except Exception as e:
        try:
            if p == 'groq':
                reply = ask_nvidia(prompt)
            elif p == 'nvidia':
                reply = ask_groq(prompt)
            else:
                reply = ask_groq(prompt)
            send_long(message.chat.id, f"[Fallback]\n\n{reply}", reply_to=message.message_id)
        except:
            bot.reply_to(message, f"Error: {str(e)}")

@app.route('/', methods=['GET'])
def index():
    return 'Telegram AI Bot is running!', 200

@app.route('/', methods=['POST'])
def webhook():
    try:
        body = request.get_data(as_text=True)
        update = telebot.types.Update.de_json(body)
        bot.process_new_updates([update])
        return ('ok', 200)
    except Exception as e:
        err = f"{type(e).__name__}: {str(e)}\n{traceback.format_exc()}"
        print(err, flush=True)
        return (err, 200)
